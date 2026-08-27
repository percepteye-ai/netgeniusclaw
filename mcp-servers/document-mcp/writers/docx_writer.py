"""Word documents. Spec 082, US1.

Provenance here is INLINE, not footnoted. python-docx 1.2.0 exposes no footnote API
(measured — research D3): `Paragraph` and `Run` have no footnote attributes at all. It
does expose `add_comment`, which is precisely the hidden mechanism FR-008a prohibits as
the means of satisfying provenance — comments are collapsed by default, stripped on
copy-paste, and absent in print.

So: tables carry a Source column the writer appends, prose figures carry a visible
parenthetical, every page footer carries the stamp, and the document ends with a
Sources section. That is more visible than a footnote, not less.
"""

from __future__ import annotations

import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from guards import apply_bound, max_blocks, reject_template, resolve_embedded_image
from outcomes import (
    Outcome,
    TaggedValue,
    ValueRefused,
    find_disagreements,
    parse_tagged,
    render_as_of,
    render_source,
    render_tagged,
)
from provenance import DocumentStamp, SourceLedger
from sanitize import plain_text

GAP_COLOR = RGBColor(0x99, 0x00, 0x00)
MUTED = RGBColor(0x60, 0x60, 0x60)

# Digits that are legitimately not a figure. Without this allow-list the prose lint
# fires on every date and ticket number and gets tuned into vacuity.
_ALLOWED_DIGIT_PATTERNS = [
    r"\d{4}-\d{2}-\d{2}(?:[T ]\d{2}:\d{2}(?::\d{2})?)?(?:Z|[+-]\d{2}:?\d{2})?",  # ISO date/time
    r"\d{1,2}:\d{2}(?::\d{2})?",                                                  # clock time
    r"[A-Z]{2,6}\d{4,}",                                                          # CHG0012345
    r"\bRFC\s?\d+\b",                                                             # RFC 6811
    r"\bv?\d+\.\d+(?:\.\d+)*\b",                                                  # 7.6.7
    r"\b(?:\d{1,3}\.){3}\d{1,3}(?:/\d{1,2})?\b",                                  # IPv4 / prefix
    r"\bAS\d+\b",                                                                 # AS65000
    r"\b(?:[0-9a-fA-F]{0,4}:){2,7}[0-9a-fA-F]{0,4}(?:/\d{1,3})?",                 # IPv6
    r"(?i:\b(?:section|table|figure|appendix|slide|step|phase|page)\s+\d+\b)",    # ordinals
]
_ALLOWED_RE = re.compile("|".join(_ALLOWED_DIGIT_PATTERNS))


def _prose_has_bare_figure(text: str) -> bool:
    """True if the prose asserts a number that is not an allow-listed pattern.

    Prose is the one place an unattributed figure could hide: `figure`, `table` and
    `keyvalue` all force attribution, and a `paragraph` does not.
    """
    return bool(re.search(r"\d", _ALLOWED_RE.sub("", text)))


def _add_source_run(paragraph, text: str) -> None:
    run = paragraph.add_run(f"  ({text})")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    run.italic = True


def _write_value_cell(cell, tv: TaggedValue) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(plain_text(render_tagged(tv)))
    if tv.is_gap:
        run.bold = True
        run.font.color.rgb = GAP_COLOR


def build(payload: dict, stamp: DocumentStamp, ledger: SourceLedger) -> tuple[bytes, list[str]]:
    """Returns (docx bytes, caveats). Raises ValueRefused for anything refusable."""
    reject_template(payload)

    blocks = payload.get("blocks") or []
    if not isinstance(blocks, list):
        raise ValueRefused("'blocks' must be a list", Outcome.REFUSED_UNTYPED)

    blocks, truncated = apply_bound(blocks, max_blocks())
    if truncated:
        stamp.truncated = True
        stamp.bound_applied = max_blocks()
        stamp.bound_kind = "blocks"

    caveats: list[str] = []
    labelled: list[tuple[str, TaggedValue]] = []

    doc = Document()
    title = plain_text(payload.get("title") or "NetClaw document")
    doc.add_heading(title, 0)

    # The stamp, visible at the top as well as in the footer. A reader who prints one
    # page still learns when this was generated.
    head = doc.add_paragraph()
    hrun = head.add_run(stamp.footer_text())
    hrun.font.size = Pt(8)
    hrun.font.color.rgb = MUTED

    if stamp.truncated:
        warn = doc.add_paragraph()
        wrun = warn.add_run(stamp.truncation_text())
        wrun.bold = True
        wrun.font.color.rgb = GAP_COLOR

    for index, block in enumerate(blocks):
        if not isinstance(block, dict):
            raise ValueRefused(f"blocks[{index}] must be an object", Outcome.REFUSED_UNTYPED)
        kind = block.get("type")
        path = f"blocks[{index}]"

        if kind == "heading":
            level = max(1, min(4, int(block.get("level", 1))))
            doc.add_heading(plain_text(block.get("text", "")), level)

        elif kind == "paragraph":
            text = plain_text(block.get("text", ""))
            doc.add_paragraph(text)
            if _prose_has_bare_figure(text):
                caveats.append(
                    f"{path}: prose asserts a figure with no source. Prose carries no "
                    f"attribution — use a 'figure', 'table' or 'keyvalue' block so the "
                    f"number is traceable"
                )

        elif kind == "figure":
            label = plain_text(block.get("label", ""))
            tv = ledger.record(parse_tagged(block.get("value"), f"{path}.value"))
            labelled.append((label, tv))
            para = doc.add_paragraph()
            para.add_run(f"{label}: ").bold = True
            vrun = para.add_run(plain_text(render_tagged(tv)))
            if tv.is_gap:
                vrun.bold = True
                vrun.font.color.rgb = GAP_COLOR
            detail = render_source(tv)
            if tv.kind == "value" and tv.as_of:
                detail += f" · as of {tv.as_of}"
            _add_source_run(para, detail)

        elif kind == "keyvalue":
            pairs = block.get("pairs") or []
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            for i, header in enumerate(["Field", "Value", "Source", "As of"]):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
            for j, pair in enumerate(pairs):
                label = plain_text(pair.get("label", ""))
                tv = ledger.record(parse_tagged(pair.get("value"), f"{path}.pairs[{j}].value"))
                labelled.append((label, tv))
                row = table.add_row().cells
                row[0].text = label
                _write_value_cell(row[1], tv)
                row[2].text = render_source(tv)
                row[3].text = render_as_of(tv)

        elif kind == "table":
            columns = list(block.get("columns") or [])
            rows = block.get("rows") or []
            if block.get("caption"):
                cap = doc.add_paragraph()
                cap.add_run(plain_text(block["caption"])).bold = True
            table = doc.add_table(rows=1, cols=len(columns) + 2)
            table.style = "Table Grid"
            for i, header in enumerate(columns + ["Source", "As of"]):
                cell = table.rows[0].cells[i]
                cell.text = plain_text(header)
                cell.paragraphs[0].runs[0].bold = True
            for r, raw_row in enumerate(rows):
                cells = table.add_row().cells
                row_tvs: list[TaggedValue] = []
                for c, raw in enumerate(list(raw_row)[: len(columns)]):
                    tv = ledger.record(parse_tagged(raw, f"{path}.rows[{r}][{c}]"))
                    row_tvs.append(tv)
                    if c < len(columns):
                        labelled.append((f"{plain_text(block.get('caption') or 'table')}."
                                         f"{plain_text(columns[c])}[{r}]", tv))
                    _write_value_cell(cells[c], tv)
                # One Source cell per row, derived from the row's own values.
                srcs = sorted({render_source(t) for t in row_tvs if t.kind == "value"})
                ages = sorted({t.as_of for t in row_tvs if t.kind == "value" and t.as_of})
                cells[len(columns)].text = " / ".join(srcs) if srcs else "— (no data)"
                cells[len(columns) + 1].text = ages[-1] if ages else "(not stated by source)"

        elif kind == "image":
            resolved = resolve_embedded_image(
                block.get("path", ""), block.get("src", ""), path
            )
            doc.add_picture(str(resolved), width=Inches(6.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            crun = cap.add_run(plain_text(block.get("caption", "")))
            crun.italic = True
            _add_source_run(cap, f"produced by {plain_text(block['src'])}")

        elif kind == "pagebreak":
            doc.add_page_break()

        else:
            raise ValueRefused(
                f"{path}: unknown block type {kind!r}. Supported: heading, paragraph, "
                f"figure, table, keyvalue, image, pagebreak",
                Outcome.REFUSED_UNTYPED,
            )

    # FR-027b — both values shown, neither dropped, no winner picked.
    for dis in find_disagreements(labelled):
        caveats.append(dis.caveat())
        para = doc.add_paragraph()
        drun = para.add_run(f"Sources disagree — {dis.label}: ")
        drun.bold = True
        drun.font.color.rgb = GAP_COLOR
        para.add_run(
            "; ".join(f"{render_tagged(v)} ({render_source(v)})" for v in dis.values)
        )
        _add_source_run(para, "not reconciled by NetClaw")

    _add_sources_section(doc, ledger)

    # Footer on every page (FR-006, FR-008).
    footer_para = doc.sections[0].footer.paragraphs[0]
    footer_para.text = stamp.footer_text()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.size = Pt(7)
        run.font.color.rgb = MUTED

    # Additive only — never the provenance mechanism (FR-008a).
    doc.core_properties.comments = stamp.footer_text()
    doc.core_properties.author = stamp.generated_by

    import io

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), caveats


def _add_sources_section(doc, ledger: SourceLedger) -> None:
    doc.add_page_break()
    doc.add_heading("Sources", 1)
    doc.add_paragraph(
        "Every figure in this document came from one of the following. This section is "
        "in addition to the per-element attribution in the tables above, not instead "
        "of it."
    )
    table = doc.add_table(rows=1, cols=len(SourceLedger.SOURCE_COLUMNS))
    table.style = "Table Grid"
    for i, header in enumerate(SourceLedger.SOURCE_COLUMNS):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    for row in ledger.as_rows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = plain_text(value)
